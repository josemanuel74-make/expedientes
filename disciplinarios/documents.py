from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from datetime import datetime

from docx import Document

from .directory import format_person_name


MARKER_RE = re.compile(r"<<([^<>]+)>>")

FIELD_ALIASES = {
    "nombrealumno": "nombreAlumno",
    "cursoalumno": "cursoAlumno",
    "grupoalumno": "grupoAlumno",
    "grupoaalumno": "grupoAlumno",
    "diahecho": "diaHechos",
    "diahechos": "diaHechos",
    "meshecho": "mesHechos",
    "meshechos": "mesHechos",
    "hechos": "hechos",
    "nombreinstructor": "nombreInstructor",
    "nombrepadres": "nombrePadres",
    "nombretutores": "nombrePadres",
    "fechaapertura": "fechaApertura",
    "mesapertura": "mesApertura",
    "cargoprimero": "cargoPrimero",
    "cargosegundo": "cargoSegundo",
    "cargotercero": "cargoTercero",
    "horasvisita": "horasVisita",
    "diavisita": "diaVisita",
    "mesvisita": "mesVisita",
    "lugarcita": "lugarCita",
    "fechahoracita": "fechaHoraCita",
    "hechosimputados": "hechosImputados",
    "calificacionhechos": "calificacionHechos",
    "propuesta": "propuesta",
    "diassuspension": "diasSuspension",
    "diasexpulsioncautelar": "diasExpulsionCautelar",
    "diaconsejoescolar": "diaConsejoEscolar",
    "mesconsejoescolar": "mesConsejoEscolar",
    "firmavisible": "firmaVisible",
}

BARE_RUN_ALLOWED_FIELDS = {
    "nombreAlumno",
    "hechos",
}

FIELD_LABELS = {
    "nombreAlumno": "Nombre del alumno",
    "cursoAlumno": "Curso",
    "grupoAlumno": "Grupo",
    "diaHechos": "Día de los hechos",
    "mesHechos": "Mes de los hechos",
    "hechos": "Hechos",
    "nombreInstructor": "Nombre del instructor",
    "nombrePadres": "Representantes legales",
    "fechaApertura": "Día de apertura",
    "mesApertura": "Mes de apertura",
    "cargoPrimero": "Cargo primero",
    "cargoSegundo": "Cargo segundo",
    "cargoTercero": "Cargo tercero",
    "horasVisita": "Hora de la vista",
    "diaVisita": "Día de la vista",
    "mesVisita": "Mes de la vista",
    "lugarCita": "Lugar de citación",
    "fechaHoraCita": "Fecha y hora de la citación",
    "hechosImputados": "Hechos imputados",
    "calificacionHechos": "Calificación de los hechos",
    "propuesta": "Propuesta",
    "diasSuspension": "Días de suspensión",
    "diasExpulsionCautelar": "Días de expulsión cautelar",
    "diaConsejoEscolar": "Día del Consejo Escolar",
    "mesConsejoEscolar": "Mes del Consejo Escolar",
    "firmaVisible": "Firma visible",
}

FIELD_HELP_TEXTS = {
    "fechaHoraCita": (
        "Escribe la fecha y la hora tal como quieres que aparezcan en el documento. "
        "Ejemplo: 27/04/2026 a las 13:30"
    ),
    "lugarCita": "Indica el lugar exacto de la comparecencia. Ejemplo: Jefatura de Estudios.",
    "horasVisita": "Escribe solo la hora. Ejemplo: 13:30",
}

MONTH_NAMES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


def slugify_marker(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_text = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]", "", ascii_text.lower())


def canonical_field(marker_name: str) -> str | None:
    return FIELD_ALIASES.get(slugify_marker(marker_name))


def bare_run_field(run_text: str) -> str | None:
    stripped = run_text.strip()
    if not stripped or "<<" in stripped or ">>" in stripped:
        return None
    field = canonical_field(stripped)
    if field not in BARE_RUN_ALLOWED_FIELDS:
        return None
    return field


def extract_markers_from_paragraph_text(text: str) -> list[str]:
    fields: list[str] = []
    for match in MARKER_RE.findall(text):
        field = canonical_field(match)
        if field and field not in fields:
            fields.append(field)
    return fields


def extract_bare_fields_from_runs(runs) -> list[str]:
    fields: list[str] = []
    for run in runs:
        field = bare_run_field(run.text)
        if field and field not in fields:
            fields.append(field)
    return fields


def template_fields(template_path: Path) -> list[str]:
    document = Document(str(template_path))
    fields: list[str] = []

    def collect_paragraph(paragraph):
        text = "".join(run.text for run in paragraph.runs)
        for field in extract_markers_from_paragraph_text(text):
            if field not in fields:
                fields.append(field)
        for field in extract_bare_fields_from_runs(paragraph.runs):
            if field not in fields:
                fields.append(field)

    for paragraph in document.paragraphs:
        collect_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_paragraph(paragraph)

    return fields


def format_marker_value(marker_name: str, data: dict[str, str]) -> str:
    field_name = canonical_field(marker_name)
    if not field_name:
        return f"<<{marker_name}>>"

    value = str(data.get(field_name, ""))
    if marker_name.isupper():
        return value.upper()
    return value


def _marker_target_run(paragraph, first_idx: int, last_idx: int) -> int:
    candidate = None
    for idx in range(first_idx, last_idx + 1):
        text = paragraph.runs[idx].text
        if re.search(r"[A-Za-zÁÉÍÓÚÑáéíóúñ]", text):
            candidate = idx
            if paragraph.runs[idx].bold:
                return idx
    return candidate if candidate is not None else first_idx


def _marker_style(involved, start_pos: int, end_pos: int, paragraph):
    marker_runs = []
    for idx, run_start, run_end in involved:
        overlap_start = max(run_start, start_pos)
        overlap_end = min(run_end, end_pos)
        if overlap_start >= overlap_end:
            continue
        marker_runs.append(paragraph.runs[idx])

    bold_values = [run.bold for run in marker_runs if run.bold is not None]
    italic_values = [run.italic for run in marker_runs if run.italic is not None]
    underline_values = [run.underline for run in marker_runs if run.underline is not None]

    return {
        "bold": True if True in bold_values else (False if False in bold_values else None),
        "italic": True if True in italic_values else (False if False in italic_values else None),
        "underline": True if True in underline_values else (False if False in underline_values else None),
    }


def replace_markers_in_paragraph(paragraph, data: dict[str, str]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if "<<" not in full_text or ">>" not in full_text:
        return

    complex_chain = any(
        re.search(r">>.*<<", run.text) for run in paragraph.runs
    )
    if complex_chain:
        replaced = MARKER_RE.sub(lambda match: format_marker_value(match.group(1), data), full_text)
        if replaced != full_text:
            paragraph.runs[0].text = replaced
            for run in paragraph.runs[1:]:
                run.text = ""
        return

    positions = []
    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        start = cursor
        end = cursor + len(run.text)
        positions.append((idx, start, end))
        cursor = end

    matches = list(MARKER_RE.finditer(full_text))
    for match in reversed(matches):
        replacement = format_marker_value(match.group(1), data)
        start_pos, end_pos = match.span()
        involved = [item for item in positions if item[1] < end_pos and item[2] > start_pos]
        if not involved:
            continue

        first_idx = involved[0][0]
        last_idx = involved[-1][0]
        target_idx = _marker_target_run(paragraph, first_idx, last_idx)
        marker_style = _marker_style(involved, start_pos, end_pos, paragraph)

        prefix = full_text[involved[0][1]:start_pos]
        suffix = full_text[end_pos:involved[-1][2]]

        paragraph.runs[target_idx].text = prefix + replacement + suffix
        if marker_style["bold"] is not None:
            paragraph.runs[target_idx].bold = marker_style["bold"]
        if marker_style["italic"] is not None:
            paragraph.runs[target_idx].italic = marker_style["italic"]
        if marker_style["underline"] is not None:
            paragraph.runs[target_idx].underline = marker_style["underline"]
        for idx in range(first_idx, last_idx + 1):
            if idx != target_idx:
                paragraph.runs[idx].text = ""


def replace_bare_tokens_in_paragraph(paragraph, data: dict[str, str]) -> None:
    for run in paragraph.runs:
        field = bare_run_field(run.text)
        if not field:
            continue

        stripped = run.text.strip()
        replacement = format_marker_value(stripped, data)
        left_padding = run.text[: len(run.text) - len(run.text.lstrip())]
        right_padding = run.text[len(run.text.rstrip()):]
        run.text = f"{left_padding}{replacement}{right_padding}"


def replace_markers_in_document(document: Document, data: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        replace_markers_in_paragraph(paragraph, data)
        replace_bare_tokens_in_paragraph(paragraph, data)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_markers_in_paragraph(paragraph, data)
                    replace_bare_tokens_in_paragraph(paragraph, data)


def template_candidates(project_root: Path) -> list[Path]:
    return sorted(project_root.glob("*.docx"))


def build_document_data(case_row, student_row) -> dict[str, str]:
    facts_date = case_row["facts_date"] or ""
    day = ""
    month = ""
    if len(facts_date) >= 10 and facts_date[4] == "-":
        _, month_num, day_num = facts_date.split("-")
        day = str(int(day_num))
        month = MONTH_NAMES.get(int(month_num), "")

    opening_date = case_row["opening_date"] or ""
    opening_day = ""
    opening_month = ""
    if len(opening_date) >= 10 and opening_date[4] == "-":
        _, opening_month_num, opening_day_num = opening_date.split("-")
        opening_day = str(int(opening_day_num))
        opening_month = MONTH_NAMES.get(int(opening_month_num), "")

    board_date = case_row["school_board_date"] or ""
    board_day = ""
    board_month = ""
    if len(board_date) >= 10 and board_date[4] == "-":
        _, board_month_num, board_day_num = board_date.split("-")
        board_day = str(int(board_day_num))
        board_month = MONTH_NAMES.get(int(board_month_num), "")

    return {
        "nombreAlumno": student_row["full_name"],
        "cursoAlumno": student_row["course_name"],
        "grupoAlumno": student_row["group_name"],
        "nombrePadres": student_row["guardians_name"],
        "diaHechos": day,
        "mesHechos": board_month if False else month,
        "hechos": case_row["facts_summary"] or "",
        "nombreInstructor": format_person_name(case_row["instructor_name"] or ""),
        "fechaApertura": opening_day,
        "mesApertura": opening_month,
        "cargoPrimero": "",
        "cargoSegundo": "",
        "cargoTercero": "",
        "horasVisita": "",
        "diaVisita": "",
        "mesVisita": "",
        "lugarCita": "",
        "fechaHoraCita": "",
        "hechosImputados": case_row["facts_summary"] or "",
        "calificacionHechos": case_row["conduct_type"] or "",
        "propuesta": case_row["proposed_measure"] or "",
        "diasSuspension": case_row["suspension_days"] or "",
        "diasExpulsionCautelar": case_row["precautionary_days"] or "",
        "diaConsejoEscolar": board_day,
        "mesConsejoEscolar": board_month,
        "firmaVisible": "",
    }


def merge_document_data(base: dict[str, str], overrides: dict[str, str]) -> dict[str, str]:
    merged = dict(base)
    for key, value in overrides.items():
        if str(value or "").strip():
            merged[key] = format_person_name(value) if key == "nombreInstructor" else value
    return merged


def month_name_from_iso(value: str) -> str:
    if not value:
        return ""
    try:
        parsed = datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        return ""
    return MONTH_NAMES.get(parsed.month, "")


def generate_document(template_path: Path, output_path: Path, data: dict[str, str]) -> None:
    document = Document(str(template_path))
    replace_markers_in_document(document, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
